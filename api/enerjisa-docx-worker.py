import base64
import importlib.util
import io
import json
from http.server import BaseHTTPRequestHandler
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from enerjisa_logo import ENERJISA_LOGO_B64

_spec = importlib.util.spec_from_file_location("jetwork_docx_base", Path(__file__).with_name("docx-worker.py"))
base = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(base)

NAVY = "17365D"
BLUE = "4472C4"
TABLE_HEADER = "D9D9D9"


def _logo():
    return io.BytesIO(base64.b64decode(ENERJISA_LOGO_B64))


def _no_borders(table):
    pr = table._tbl.tblPr
    borders = pr.first_child_found_in("w:tblBorders") or OxmlElement("w:tblBorders")
    if borders.getparent() is None:
        pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        if node.getparent() is None:
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _cell_border(cell, edge, color="808080", size="12"):
    pr = cell._tc.get_or_add_tcPr()
    borders = pr.first_child_found_in("w:tcBorders") or OxmlElement("w:tcBorders")
    if borders.getparent() is None:
        pr.append(borders)
    node = borders.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
    if node.getparent() is None:
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), size)
    node.set(qn("w:color"), color)


def _page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def _meta(metadata, label):
    wanted = label.casefold()
    for item in metadata:
        if isinstance(item, dict) and base._clean(item.get("label"), 200).casefold() == wanted:
            return base._clean(item.get("value"), 2000)
    return ""


def _styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(4)
    for name, size in (("Title", 26), ("Heading 1", 15), ("Heading 2", 11), ("Heading 3", 9.5)):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)


def _header_footer(section):
    section.different_first_page_header_footer = True
    header = section.header
    table = header.add_table(1, 2, Inches(7.0))
    _no_borders(table)
    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    r = p.add_run("İş Analizi Dokümanı\nİş Uygulamaları Yönetim Müdürlüğü\nEnerjisa Elektrik Perakende Satış A.Ş")
    r.font.name = "Arial"
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run().add_picture(_logo(), width=Inches(1.45))
    for footer in (section.footer, section.first_page_footer):
        ft = footer.add_table(1, 3, Inches(7.0))
        _no_borders(ft)
        center = ft.rows[0].cells[1].paragraphs[0]
        center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _page_number(center)
        rightp = ft.rows[0].cells[2].paragraphs[0]
        rightp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rightp.add_run("Hizmete Özel\nGizli")
        rr.font.name = "Arial"
        rr.font.size = Pt(7.5)


def _cover(doc, title, metadata):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(_logo(), width=Inches(1.75))
    for _ in range(6):
        doc.add_paragraph()
    table = doc.add_table(1, 2)
    _no_borders(table)
    left, right = table.rows[0].cells
    _cell_border(left, "bottom")
    _cell_border(left, "right")
    _cell_border(right, "bottom")
    lp = left.paragraphs[0]
    lr = lp.add_run("İş Analizi\nDokümanı")
    lr.font.name = "Arial"
    lr.font.size = Pt(27)
    lr.font.color.rgb = RGBColor.from_string(NAVY)
    lr.font.underline = True
    rp = right.paragraphs[0]
    rr = rp.add_run("Talep Adı:\n")
    rr.bold = True
    rr.font.name = "Arial"
    rr.font.size = Pt(14)
    rr.font.color.rgb = RGBColor.from_string(NAVY)
    rv = rp.add_run(_meta(metadata, "Talep Adı") or title or "[AÇIK KONU]")
    rv.font.name = "Arial"
    rv.font.size = Pt(12)
    rv.font.color.rgb = RGBColor.from_string(NAVY)
    no = doc.add_paragraph()
    rn = no.add_run("Talep No: ")
    rn.bold = True
    rn.font.color.rgb = RGBColor.from_string(NAVY)
    no.add_run(_meta(metadata, "Talep No") or "[AÇIK KONU]")
    doc.add_page_break()


def _brand_body(doc):
    for table in doc.tables[1:]:
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            base._shade(cell, TABLE_HEADER)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Arial"
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(NAVY)
        for row in table.rows[1:]:
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
    for p in doc.paragraphs:
        text = p.text.strip().upper()
        if text == "İHTİYAÇ ANALİZİ":
            p.style = doc.styles["Heading 1"]
            if p.runs:
                p.runs[0].font.size = Pt(17)
            ppr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), BLUE)
            pbdr.append(bottom)
            ppr.append(pbdr)


def _build_docx(payload):
    title = base._clean(payload.get("title"), 500)
    paragraphs = payload.get("paragraphs") if isinstance(payload.get("paragraphs"), list) else []
    markdown = payload.get("markdown")
    if not isinstance(markdown, str) or not markdown.strip():
        markdown = "\n\n".join(str(v)[:8000] for v in paragraphs[:500] if str(v).strip())
    markdown = markdown[:base.MAX_MARKDOWN_CHARS]
    if not title and not markdown.strip():
        raise ValueError("DOCX content is empty.")
    metadata = payload.get("metadata") if isinstance(payload.get("metadata"), list) else []
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.60)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    _styles(doc)
    _header_footer(section)
    _cover(doc, title, metadata)
    stats = base._render_markdown(doc, markdown)
    _brand_body(doc)
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    if len(raw) < 1000 or not raw.startswith(b"PK"):
        raise RuntimeError("python-docx generated an invalid package.")
    reloaded = Document(io.BytesIO(raw))
    return raw, {
        "engine": "python-docx-enerjisa",
        "brandProfile": "enerjisa_analysis",
        "logoEmbedded": True,
        "packageReloaded": True,
        "paragraphCount": len(reloaded.paragraphs),
        "tableCount": len(reloaded.tables),
        **stats,
    }


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            length = int(self.headers.get("Content-Length") or "0")
            if length <= 0 or length > base.MAX_REQUEST_BYTES:
                return base._json_response(self, 413, {"error": "DOCX worker request is empty or too large."})
            authorization = self.headers.get("Authorization") or ""
            base._verify_supabase_user(authorization)
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
            if not isinstance(payload, dict):
                return base._json_response(self, 400, {"error": "DOCX worker request must be a JSON object."})
            raw, qa = _build_docx(payload)
            return base._json_response(self, 200, {
                "mimeType": base.DOCX_MIME,
                "bytesBase64": base64.b64encode(raw).decode("ascii"),
                "byteSize": len(raw),
                "qa": qa,
            })
        except PermissionError as error:
            return base._json_response(self, 401, {"error": str(error)})
        except ValueError as error:
            return base._json_response(self, 400, {"error": str(error)})
        except Exception as error:
            return base._json_response(self, 500, {"error": f"Enerjisa DOCX generation failed: {str(error)[:1000]}"})

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Headers", "authorization, content-type")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.end_headers()
