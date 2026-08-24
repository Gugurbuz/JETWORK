import base64, io, json, os, re, urllib.error, urllib.request
from http.server import BaseHTTPRequestHandler
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from enerjisa_brand_asset import ENERJISA_LOGO_BASE64
except ImportError:
    from api.enerjisa_brand_asset import ENERJISA_LOGO_BASE64

DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
MAX_REQUEST_BYTES = 1_600_000
MAX_MARKDOWN_CHARS = 400_000
DARK = '0F243E'; ACCENT = '4F81BD'; HEADER = '365F91'; TABLE = 'D9D9D9'; LIGHT = 'F2F2F2'
TABLE_SEPARATOR_RE = re.compile(r'^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$')
HEADING_RE = re.compile(r'^(#{1,6})\s+(.+?)\s*$')
BULLET_RE = re.compile(r'^\s*[-*+]\s+(.+)$')
NUMBER_RE = re.compile(r'^\s*\d+[.)]\s+(.+)$')


def clean(value, limit=500): return str(value or '').strip()[:limit]

def json_response(h, status, payload):
    raw = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    h.send_response(status); h.send_header('Content-Type','application/json; charset=utf-8'); h.send_header('Content-Length',str(len(raw))); h.end_headers(); h.wfile.write(raw)

def verify_user(auth):
    if not auth or not auth.lower().startswith('bearer '): raise ValueError('Authorization bearer token is required.')
    url = os.environ.get('VITE_SUPABASE_URL') or os.environ.get('SUPABASE_URL')
    key = os.environ.get('VITE_SUPABASE_ANON_KEY') or os.environ.get('SUPABASE_ANON_KEY')
    if not url or not key: raise RuntimeError('Supabase auth configuration is missing in the DOCX worker.')
    req = urllib.request.Request(f"{url.rstrip('/')}/auth/v1/user", headers={'Authorization':auth,'apikey':key}, method='GET')
    try:
        with urllib.request.urlopen(req, timeout=8) as r: payload=json.loads(r.read().decode('utf-8'))
    except urllib.error.HTTPError as e: raise PermissionError(f'Supabase user verification failed ({e.code}).') from e
    if not clean(payload.get('id'),160): raise PermissionError('Supabase user verification returned no user id.')

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); node=pr.find(qn('w:shd'))
    if node is None: node=OxmlElement('w:shd'); pr.append(node)
    node.set(qn('w:fill'),fill)

def margins(cell, value=70):
    pr=cell._tc.get_or_add_tcPr(); tc=pr.first_child_found_in('w:tcMar')
    if tc is None: tc=OxmlElement('w:tcMar'); pr.append(tc)
    for name in ('top','start','bottom','end'):
        n=tc.find(qn(f'w:{name}'))
        if n is None: n=OxmlElement(f'w:{name}'); tc.append(n)
        n.set(qn('w:w'),str(value)); n.set(qn('w:type'),'dxa')

def cell_border(cell, color='7F8C9A', size='5', sides=('top','left','bottom','right')):
    pr=cell._tc.get_or_add_tcPr(); b=pr.first_child_found_in('w:tcBorders')
    if b is None: b=OxmlElement('w:tcBorders'); pr.append(b)
    for side in sides:
        n=b.find(qn(f'w:{side}'))
        if n is None: n=OxmlElement(f'w:{side}'); b.append(n)
        n.set(qn('w:val'),'single'); n.set(qn('w:sz'),size); n.set(qn('w:space'),'0'); n.set(qn('w:color'),color)

def borderless(table):
    pr=table._tbl.tblPr; b=pr.first_child_found_in('w:tblBorders')
    if b is None: b=OxmlElement('w:tblBorders'); pr.append(b)
    for side in ('top','left','bottom','right','insideH','insideV'):
        n=b.find(qn(f'w:{side}'))
        if n is None: n=OxmlElement(f'w:{side}'); b.append(n)
        n.set(qn('w:val'),'nil')

def paragraph_bottom(paragraph):
    pr=paragraph._p.get_or_add_pPr(); b=pr.find(qn('w:pBdr'))
    if b is None: b=OxmlElement('w:pBdr'); pr.append(b)
    n=OxmlElement('w:bottom'); n.set(qn('w:val'),'single'); n.set(qn('w:sz'),'8'); n.set(qn('w:space'),'4'); n.set(qn('w:color'),ACCENT); b.append(n)

def logo_stream(): return io.BytesIO(base64.b64decode(ENERJISA_LOGO_BASE64))

def configure_styles(doc):
    normal=doc.styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(9); normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.05
    for name,size in (('Title',16),('Heading 1',13.5),('Heading 2',11.5),('Heading 3',10)):
        s=doc.styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(DARK); s.font.bold=True; s.paragraph_format.space_before=Pt(8 if name!='Title' else 0); s.paragraph_format.space_after=Pt(5)

def add_page_number(p):
    r=p.add_run(); begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin'); instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '; sep=OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate'); t=OxmlElement('w:t'); t.text='1'; end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end'); r._r.extend([begin,instr,sep,t,end])

def configure_shell(section):
    section.different_first_page_header_footer=True; section.header_distance=Inches(.18); section.footer_distance=Inches(.18)
    h=section.header; h.paragraphs[0].clear(); table=h.add_table(rows=1,cols=2,width=Inches(7.3)); borderless(table); left,right=table.rows[0].cells
    lp=left.paragraphs[0]
    for i,line in enumerate(('İş Analizi Dokümanı','İş Uygulamaları Yönetim Müdürlüğü','Enerjisa Elektrik Perakende Satış A.Ş')):
        r=lp.add_run(('\n' if i else '')+line); r.font.name='Calibri'; r.font.size=Pt(7.5); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(DARK)
    rp=right.paragraphs[0]; rp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; rp.add_run().add_picture(logo_stream(),width=Inches(1.45))
    section.first_page_header.paragraphs[0].clear()
    f=section.footer; f.paragraphs[0].clear(); ft=f.add_table(rows=1,cols=3,width=Inches(7.3)); borderless(ft); _,c,r=ft.rows[0].cells; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; add_page_number(c.paragraphs[0]); rr=r.paragraphs[0]; rr.alignment=WD_ALIGN_PARAGRAPH.RIGHT; x=rr.add_run('Hizmete Özel\nGizli'); x.font.name='Calibri'; x.font.size=Pt(8)
    ff=section.first_page_footer; ff.paragraphs[0].clear(); fft=ff.add_table(rows=1,cols=2,width=Inches(7.3)); borderless(fft); rr=fft.rows[0].cells[1].paragraphs[0]; rr.alignment=WD_ALIGN_PARAGRAPH.RIGHT; x=rr.add_run('Gizli\nHizmete Özel'); x.font.name='Calibri'; x.font.size=Pt(8)

def metadata_map(payload):
    out={}
    for item in payload.get('metadata') if isinstance(payload.get('metadata'),list) else []:
        if isinstance(item,dict) and clean(item.get('label')): out[clean(item.get('label'),200).casefold()]=clean(item.get('value'),2000)
    return out

def add_cover(doc,payload):
    meta=metadata_map(payload); title=clean(payload.get('title'),500); name=meta.get('talep adı') or meta.get('talep adi') or title or '[AÇIK KONU]'; no=meta.get('talep no') or meta.get('talep numarası') or meta.get('talep numarasi') or '[AÇIK KONU]'
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(125); p.add_run().add_picture(logo_stream(),width=Inches(1.6))
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; borderless(t); left,right=t.rows[0].cells
    for cell in (left,right): margins(cell,100); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell_border(cell,'808080','12',('bottom',))
    cell_border(left,'808080','12',('right',)); lp=left.paragraphs[0]; r=lp.add_run('İş Analizi\nDokümanı'); r.font.name='Calibri'; r.font.size=Pt(27); r.font.color.rgb=RGBColor.from_string(DARK); r.font.underline=True
    rp=right.paragraphs[0]; a=rp.add_run('Talep Adı:\n'); a.bold=True; a.font.name='Calibri'; a.font.size=Pt(12); a.font.color.rgb=RGBColor.from_string(DARK); b=rp.add_run(name); b.font.name='Calibri'; b.font.size=Pt(11)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(36); r=p.add_run(f'Talep No: {no}'); r.font.name='Calibri'; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(DARK); doc.add_page_break()

def inline(p,text):
    token=re.compile(r'(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)'); pos=0
    for m in token.finditer(text):
        if m.start()>pos: p.add_run(text[pos:m.start()])
        v=m.group(0); r=p.add_run(v[2:-2] if v.startswith('**') else v[1:-1]); r.bold=v.startswith('**'); r.italic=v.startswith('*') and not v.startswith('**');
        if v.startswith('`'): r.font.name='Consolas'; r.font.size=Pt(8.5)
        pos=m.end()
    if pos<len(text): p.add_run(text[pos:])
    for r in p.runs:
        if r.font.name!='Consolas': r.font.name='Calibri'
        if not r.font.size: r.font.size=Pt(9)

def split_row(line): return [x.strip().replace('\\|','|') for x in line.strip().strip('|').split('|')]

def add_table(doc,rows):
    cols=max(len(r) for r in rows); t=doc.add_table(rows=0,cols=cols); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; scope=len(rows[0])>=2 and rows[0][0].strip().casefold()=='başlık' and rows[0][1].strip().casefold()=='açıklama'
    for ri,values in enumerate(rows):
        cells=t.add_row().cells
        for ci,cell in enumerate(cells):
            margins(cell); cell_border(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; inline(cell.paragraphs[0], values[ci] if ci<len(values) else '')
            if ri==0: shade(cell,TABLE); [setattr(run,'bold',True) for run in cell.paragraphs[0].runs]
            elif scope and ci==0: shade(cell,LIGHT); [setattr(run,'bold',True) for run in cell.paragraphs[0].runs]
    doc.add_paragraph()

def strip_old_cover(md):
    lines=md.replace('\r\n','\n').replace('\r','\n').split('\n'); first=next((i for i,x in enumerate(lines) if x.strip()),None)
    if first is None: return md
    sample='\n'.join(lines[first:first+5]).casefold()
    if 'iş analizi dokümanı' not in sample or 'talep adı' not in sample: return md
    i=first
    while i<len(lines) and (not lines[i].strip() or '|' in lines[i]): i+=1
    return '\n'.join(lines[:first]+lines[i:]).lstrip()

def render_markdown(doc,md):
    lines=strip_old_cover(md).replace('\r\n','\n').replace('\r','\n').split('\n'); i=0; hc=tc=0
    while i<len(lines):
        s=lines[i].strip()
        if not s: i+=1; continue
        if '|' in lines[i] and i+1<len(lines) and TABLE_SEPARATOR_RE.match(lines[i+1]):
            rows=[split_row(lines[i])]; i+=2
            while i<len(lines) and lines[i].strip() and '|' in lines[i]: rows.append(split_row(lines[i])); i+=1
            add_table(doc,rows); tc+=1; continue
        h=HEADING_RE.match(lines[i])
        if h:
            level=min(len(h.group(1)),3); text=re.sub(r'\s+#+\s*$','',h.group(2)).strip(); p=doc.add_heading(text,level=level)
            if text.upper()=='İHTİYAÇ ANALİZİ': p.style=doc.styles['Title']; paragraph_bottom(p)
            hc+=1; i+=1; continue
        m=BULLET_RE.match(lines[i]) or NUMBER_RE.match(lines[i])
        if m: p=doc.add_paragraph(style='List Bullet' if BULLET_RE.match(lines[i]) else 'List Number'); inline(p,m.group(1)); i+=1; continue
        if s in ('---','***','___'): i+=1; continue
        p=doc.add_paragraph(); inline(p,s);
        if (s.startswith('[AÇIK KONU]') or s.startswith('[VARSAYIM]')) and p.runs: p.runs[0].bold=True
        i+=1
    return {'headingCount':hc,'tableCount':tc}

def build_docx(payload):
    title=clean(payload.get('title'),500); paragraphs=payload.get('paragraphs') if isinstance(payload.get('paragraphs'),list) else []; md=payload.get('markdown')
    if not isinstance(md,str) or not md.strip(): md='\n\n'.join(str(v)[:8000] for v in paragraphs[:500] if str(v).strip())
    md=md[:MAX_MARKDOWN_CHARS]
    if not title and not md.strip(): raise ValueError('DOCX content is empty.')
    doc=Document(); s=doc.sections[0]; s.top_margin=Inches(.79); s.bottom_margin=Inches(.30); s.left_margin=Inches(.50); s.right_margin=Inches(.40); configure_styles(doc); configure_shell(s); add_cover(doc,payload); stats=render_markdown(doc,md)
    buf=io.BytesIO(); doc.save(buf); raw=buf.getvalue()
    if len(raw)<1000 or not raw.startswith(b'PK'): raise RuntimeError('python-docx generated an invalid package.')
    reloaded=Document(io.BytesIO(raw)); qa={'engine':'python-docx','brandProfile':'enerjisa-analysis-v1','brandLogoEmbedded':True,'packageReloaded':True,'paragraphCount':len(reloaded.paragraphs),'tableCount':len(reloaded.tables),**stats}; return raw,qa

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            n=int(self.headers.get('Content-Length') or '0')
            if n<=0 or n>MAX_REQUEST_BYTES: return json_response(self,413,{'error':'DOCX worker request is empty or too large.'})
            verify_user(self.headers.get('Authorization') or ''); payload=json.loads(self.rfile.read(n).decode('utf-8'))
            if not isinstance(payload,dict): return json_response(self,400,{'error':'DOCX worker request must be a JSON object.'})
            raw,qa=build_docx(payload); return json_response(self,200,{'mimeType':DOCX_MIME,'bytesBase64':base64.b64encode(raw).decode('ascii'),'byteSize':len(raw),'qa':qa})
        except PermissionError as e: return json_response(self,401,{'error':str(e)})
        except ValueError as e: return json_response(self,400,{'error':str(e)})
        except Exception as e: return json_response(self,500,{'error':f'DOCX generation failed: {str(e)[:1000]}'})
    def do_OPTIONS(self): self.send_response(204); self.send_header('Access-Control-Allow-Origin','*'); self.send_header('Access-Control-Allow-Headers','authorization, content-type'); self.send_header('Access-Control-Allow-Methods','POST, OPTIONS'); self.end_headers()
    def log_message(self, format, *args): return
