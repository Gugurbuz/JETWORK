import base64
import io
import json
import os
import re
import urllib.error
import urllib.request
from http.server import BaseHTTPRequestHandler

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_REQUEST_BYTES = 1_600_000
MAX_MARKDOWN_CHARS = 400_000
NAVY = "17365D"
BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
GRAY = "666666"
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
NUMBER_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")


def _clean(value, limit=500):
    return str(value or "").strip()[:limit]


def _json_response(handler, status, payload):
    raw = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(raw)))
    handler.end_headers()
    handler.wfile.write(raw)


def _verify_supabase_user(authorization):
    if not authorization or not authorization.lower().startswith("bearer "):
        raise ValueError("Authorization bearer token is required.")
    supabase_url = os.environ.get("VITE_SUPABASE_URL") or os.environ.get("SUPABASE_URL")
    anon_key = os.environ.get("VITE_SUPABASE_ANON_KEY") or os.environ.get("SUPABASE_ANON_KEY")
    if not supabase_url or not anon_key:
        raise RuntimeError("Supabase auth configuration is missing in the DOCX worker.")
    request = urllib.request.Request(
        f"{supabase_url.rstrip('/')}/auth/v1/user",
        headers={"Authorization": authorization, "apikey": anon_key},
        method="GET",
    )
    try:
        with urllib.request.urlopen(request, timeout=8) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        raise PermissionError(f"Supabase user verification failed ({error.code}).") from error
    user_id = _clean(payload.get("id"), 160)
    if not user_id:
        raise PermissionError("Supabase user verification returned no user id.")
    return user_id


def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell, margin=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def _configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 10.5, BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name in ("Title", "Heading 1") else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def _add_inline_runs(paragraph, text):
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _split_table_row(line):
    stripped = line.strip().strip("|")
    return [cell.strip().replace("\\|", "|") for cell in stripped.split("|")]


def _add_markdown_table(doc, rows):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index in range(column_count):
            cell = cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _cell_margins(cell)
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            _add_inline_runs(paragraph, value)
            for run in paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(8)
            if row_index == 0:
                _shade(cell, NAVY)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
    doc.add_paragraph()


def _add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)


def _render_markdown(doc, markdown):
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    heading_count = 0
    table_count = 0
    code_fence = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if code_fence:
                _add_code_block(doc, code_lines)
                code_lines = []
                code_fence = False
            else:
                code_fence = True
            i += 1
            continue
        if code_fence:
            code_lines.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if (
            "|" in line
            and i + 1 < len(lines)
            and TABLE_SEPARATOR_RE.match(lines[i + 1])
        ):
            table_rows = [_split_table_row(line)]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_rows.append(_split_table_row(lines[i]))
                i += 1
            _add_markdown_table(doc, table_rows)
            table_count += 1
            continue
        heading = HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 3)
            doc.add_heading(re.sub(r"\s+#+\s*$", "", heading.group(2)).strip(), level=level)
            heading_count += 1
            i += 1
            continue
        bullet = BULLET_RE.match(line)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(paragraph, bullet.group(1))
            i += 1
            continue
        numbered = NUMBER_RE.match(line)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            _add_inline_runs(paragraph, numbered.group(1))
            i += 1
            continue
        if stripped.startswith(">"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(stripped.lstrip("> "))
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(GRAY)
            i += 1
            continue
        if stripped in ("---", "***", "___"):
            i += 1
            continue
        paragraph = doc.add_paragraph()
        _add_inline_runs(paragraph, stripped)
        if stripped.startswith("[AÇIK KONU]") or stripped.startswith("[VARSAYIM]"):
            paragraph.runs[0].bold = True if paragraph.runs else False
        i += 1
    if code_fence and code_lines:
        _add_code_block(doc, code_lines)
    return {"headingCount": heading_count, "tableCount": table_count}


def _add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def _build_docx(payload):
    title = _clean(payload.get("title"), 500)
    paragraphs = payload.get("paragraphs") if isinstance(payload.get("paragraphs"), list) else []
    markdown = payload.get("markdown")
    if not isinstance(markdown, str) or not markdown.strip():
        markdown = "\n\n".join(str(value)[:8000] for value in paragraphs[:500] if str(value).strip())
    markdown = markdown[:MAX_MARKDOWN_CHARS]
    if not title and not markdown.strip():
        raise ValueError("DOCX content is empty.")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    _configure_styles(doc)

    header_text = _clean(payload.get("headerText"), 500) or "JETWORK | İŞ ANALİZİ"
    footer_text = _clean(payload.get("footerText"), 500) or "JetWork tarafından oluşturuldu"
    header = section.header.paragraphs[0]
    header.text = header_text
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(BLUE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(footer_text + " | ")
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(GRAY)
    _add_page_number(footer)

    if title:
        title_paragraph = doc.add_paragraph(style="Title")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.add_run(title)
        metadata = payload.get("metadata") if isinstance(payload.get("metadata"), list) else []
        if metadata:
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for item in metadata[:20]:
                if not isinstance(item, dict):
                    continue
                label = _clean(item.get("label"), 200)
                value = _clean(item.get("value"), 2000)
                if not label and not value:
                    continue
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = value
                _shade(cells[0], LIGHT_BLUE)
                cells[0].paragraphs[0].runs[0].bold = True
                for cell in cells:
                    _cell_margins(cell)
            doc.add_paragraph()
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note.add_run("Doğrulanmamış teknik ayrıntılar varsayım veya açık konu olarak işaretlenmelidir.")
        note_run.italic = True
        note_run.font.size = Pt(8.5)
        note_run.font.color.rgb = RGBColor.from_string(GRAY)
        doc.add_page_break()

    stats = _render_markdown(doc, markdown)
    buffer = io.BytesIO()
    doc.save(buffer)
    raw = buffer.getvalue()
    if len(raw) < 1000 or not raw.startswith(b"PK"):
        raise RuntimeError("python-docx generated an invalid package.")
    reloaded = Document(io.BytesIO(raw))
    qa = {
        "engine": "python-docx",
        "packageReloaded": True,
        "paragraphCount": len(reloaded.paragraphs),
        "tableCount": len(reloaded.tables),
        **stats,
    }
    return raw, qa


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers.get("Content-Length") or "0")
            if content_length <= 0 or content_length > MAX_REQUEST_BYTES:
                return _json_response(self, 413, {"error": "DOCX worker request is empty or too large."})
            authorization = self.headers.get("Authorization") or ""
            _verify_supabase_user(authorization)
            payload = json.loads(self.rfile.read(content_length).decode("utf-8"))
            if not isinstance(payload, dict):
                return _json_response(self, 400, {"error": "DOCX worker request must be a JSON object."})
            raw, qa = _build_docx(payload)
            return _json_response(
                self,
                200,
                {
                    "mimeType": DOCX_MIME,
                    "bytesBase64": base64.b64encode(raw).decode("ascii"),
                    "byteSize": len(raw),
                    "qa": qa,
                },
            )
        except PermissionError as error:
            return _json_response(self, 401, {"error": str(error)})
        except ValueError as error:
            return _json_response(self, 400, {"error": str(error)})
        except Exception as error:
            return _json_response(self, 500, {"error": f"DOCX generation failed: {str(error)[:1000]}"})

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Allow", "POST, OPTIONS")
        self.end_headers()
